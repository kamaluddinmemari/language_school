"""
سازنده‌های خروجی: اکسل (openpyxl)، ورد قابل‌ویرایش (python-docx)، و HTML چاپی — همه راست‌به‌چپ.
"""
import io
from django.utils import timezone
import jdatetime


def _now_jalali_str():
    return jdatetime.datetime.fromgregorian(datetime=timezone.localtime(timezone.now())).strftime('%Y/%m/%d - %H:%M')


# ---------------------------------------------------------------------------
# اکسل
# ---------------------------------------------------------------------------
def build_excel(spec, rows, totals):
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    wb = Workbook()
    ws = wb.active
    ws.title = spec.name[:31] or 'گزارش'
    ws.sheet_view.rightToLeft = True

    header_font = Font(bold=True, color='FFFFFF', size=11)
    header_fill = PatternFill(start_color='1A1A2E', end_color='1A1A2E', fill_type='solid')
    thin = Side(style='thin', color='DDDDDD')
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    center = Alignment(horizontal='center', vertical='center', wrap_text=True)

    title_cell = ws.cell(row=1, column=1, value=spec.name)
    title_cell.font = Font(bold=True, size=14)
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=max(len(spec.columns), 1))
    title_cell.alignment = center

    meta_cell = ws.cell(row=2, column=1, value=f"تاریخ تهیه گزارش: {_now_jalali_str()}    |    تعداد ردیف: {len(rows)}")
    ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=max(len(spec.columns), 1))

    header_row = 4
    for idx, col in enumerate(spec.columns, start=1):
        cell = ws.cell(row=header_row, column=idx, value=col.header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = center
        cell.border = border

    r = header_row + 1
    for row in rows:
        for idx, col in enumerate(spec.columns, start=1):
            cell = ws.cell(row=r, column=idx, value=row.get(col.key, ''))
            cell.border = border
            cell.alignment = Alignment(horizontal='center', vertical='center')
        r += 1

    if totals:
        r += 1
        label_cell = ws.cell(row=r, column=1, value=spec.totals_row_label)
        label_cell.font = Font(bold=True)
        for idx, col in enumerate(spec.columns, start=1):
            if col.key in totals:
                cell = ws.cell(row=r, column=idx, value=totals[col.key])
                cell.font = Font(bold=True)
                cell.border = border

    for idx, col in enumerate(spec.columns, start=1):
        max_len = max([len(str(col.header))] + [len(str(row.get(col.key, ''))) for row in rows]) if rows else len(col.header)
        ws.column_dimensions[chr(64 + idx) if idx <= 26 else 'A'].width = min(max(max_len + 4, 12), 40)

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf


# ---------------------------------------------------------------------------
# ورد قابل‌ویرایش
# ---------------------------------------------------------------------------
def _set_rtl_paragraph(paragraph):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.append(bidi)
    for run in paragraph.runs:
        rPr = run._r.get_or_add_rPr()
        rtl = OxmlElement('w:rtl')
        rPr.append(rtl)


def _set_rtl_table(table):
    from docx.oxml import OxmlElement
    tblPr = table._tbl.tblPr
    bidi = OxmlElement('w:bidiVisual')
    tblPr.append(bidi)


def build_docx(spec, rows, totals):
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    for section in doc.sections:
        section.right_margin = section.right_margin
    for style_name in ['Normal']:
        style = doc.styles[style_name]
        style.font.name = 'B Nazanin'
        style.font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(spec.name)
    run.bold = True
    run.font.size = Pt(16)
    _set_rtl_paragraph(title)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"تاریخ تهیه گزارش: {_now_jalali_str()}    |    تعداد ردیف: {len(rows)}")
    _set_rtl_paragraph(meta)

    table = doc.add_table(rows=1, cols=len(spec.columns))
    table.style = 'Light Grid Accent 1'
    _set_rtl_table(table)
    hdr_cells = table.rows[0].cells
    for idx, col in enumerate(spec.columns):
        hdr_cells[idx].text = col.header
        for p in hdr_cells[idx].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
            _set_rtl_paragraph(p)

    for row in rows:
        cells = table.add_row().cells
        for idx, col in enumerate(spec.columns):
            cells[idx].text = str(row.get(col.key, ''))
            for p in cells[idx].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _set_rtl_paragraph(p)

    if totals:
        trow = table.add_row().cells
        trow[0].text = spec.totals_row_label
        for p in trow[0].paragraphs:
            for r in p.runs:
                r.bold = True
            _set_rtl_paragraph(p)
        for idx, col in enumerate(spec.columns):
            if col.key in totals:
                trow[idx].text = str(totals[col.key])
                for p in trow[idx].paragraphs:
                    for r in p.runs:
                        r.bold = True
                    _set_rtl_paragraph(p)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ---------------------------------------------------------------------------
# HTML چاپی
# ---------------------------------------------------------------------------
def build_print_html(spec, rows, totals):
    head_cells = ''.join(f"<th>{c.header}</th>" for c in spec.columns)
    body_rows = ''
    for row in rows:
        cells = ''.join(f"<td>{row.get(c.key, '')}</td>" for c in spec.columns)
        body_rows += f"<tr>{cells}</tr>"
    totals_row = ''
    if totals:
        cells = []
        for idx, c in enumerate(spec.columns):
            if idx == 0:
                cells.append(f"<td><b>{spec.totals_row_label}</b></td>")
            elif c.key in totals:
                cells.append(f"<td><b>{totals[c.key]}</b></td>")
            else:
                cells.append("<td></td>")
        totals_row = f"<tr class='totals'>{''.join(cells)}</tr>"

    return f"""<!DOCTYPE html>
<html dir="rtl" lang="fa">
<head>
<meta charset="utf-8">
<title>{spec.name}</title>
<style>
  body {{ font-family: Tahoma, Vazirmatn, sans-serif; direction: rtl; padding: 24px; color: #1a1a2e; }}
  h1 {{ text-align: center; font-size: 20px; margin-bottom: 4px; }}
  .meta {{ text-align: center; color: #666; font-size: 12px; margin-bottom: 20px; }}
  table {{ width: 100%; border-collapse: collapse; font-size: 13px; }}
  th, td {{ border: 1px solid #ccc; padding: 6px 8px; text-align: center; }}
  th {{ background: #1a1a2e; color: #fff; }}
  tr:nth-child(even) {{ background: #f7f7fb; }}
  .totals td {{ background: #eef1ff; }}
  @media print {{ body {{ padding: 0; }} }}
</style>
</head>
<body>
  <h1>{spec.name}</h1>
  <div class="meta">تاریخ تهیه گزارش: {_now_jalali_str()} &nbsp;|&nbsp; تعداد ردیف: {len(rows)}</div>
  <table>
    <thead><tr>{head_cells}</tr></thead>
    <tbody>{body_rows}{totals_row}</tbody>
  </table>
  <script>window.onload = function() {{ window.print(); }};</script>
</body>
</html>"""
